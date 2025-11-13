import pandas as pd
import numpy as np
import textwrap
import re
import os
import time
import asyncio
import PyPDF2
import docx  
from langchain_openai import ChatOpenAI
from typing import List, Dict, Optional, TypedDict, Any
from openai import OpenAI
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from pypdf import PdfReader
from langchain_core.runnables import RunnableParallel
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters.character import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader
from langchain_core.output_parsers import StrOutputParser
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langgraph.graph import StateGraph, END

# Get the directory of the currently executing file
BACKEND_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# Embedding model configuration
MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
K = 3 # Number of top results to retrieve
embeddings = HuggingFaceEmbeddings(model_name=MODEL_NAME)

# LLM configuration for openai
LLM_MODEL = "gpt-4o-mini" 
MAX_RETRIES = 3
load_dotenv()  # Load environment variables
api_key_from_env = os.getenv("OPENAI_API_KEY") # Get the API key from the environment variable

# Pass the key explicitly to ensure the client is initialized correctly
if api_key_from_env:
    try:
        client = OpenAI(api_key=api_key_from_env)
        print("OpenAI client initialized successfully from environment variable.")
    except Exception as e:
        print(f"FATAL ERROR: Failed to initialize OpenAI client even with key. Error: {e}")
        client = None
else:
    print("OPENAI_API_KEY environment variable is NOT set.")
    client = None

# Building RAG 1: Ideal Clauses
def build_ideal_clauses_retriever(data_directory="./TA_template", faiss_index_path="./faiss_index_ideal_clauses"):
    print("\n--- BUILDING RAG 1: IDEAL CLAUSES RETRIEVER ---")
    
    # 1. Load Documents
    loader = DirectoryLoader(
        path=data_directory,
        glob="**/*.pdf",
        loader_cls=PyPDFLoader,
        show_progress=True
    )
    all_documents = loader.load()
    print(f"Loaded {len(all_documents)} document pages.")

    # 2. Chunk Documents (using the custom, structure-aware splitter)
    custom_separators = [
        "\n\n",
        r"\n\s*[A-Z]+\s+\d*\s*\.",
        r"\n\s*\d+\.\d*\s*",
        r"\n\s*\([a-zA-Z0-9]+\)\s*",
        "\n", " ", ""
    ]
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, 
        chunk_overlap=200,
        separators=custom_separators,
        is_separator_regex=True
    )
    all_chunks = text_splitter.split_documents(all_documents)
    print(f"Split into {len(all_chunks)} chunks.")

    print("\n--- IDEAL CLAUSES CHUNK PREVIEW (First 3) ---")
    for i, chunk in enumerate(all_chunks[:3]):
        print(f"Chunk {i+1} (Length: {len(chunk.page_content)}):")
        # Print the first 200 characters to keep the output manageable
        preview = chunk.page_content[:200].replace("\n", " ")
        print(f"  {preview}...")
        print(f"  Metadata: {chunk.metadata}")
        print("-" * 20)


    # 3. Create Vector Store and Retriever
    vectorstore = FAISS.from_documents(all_chunks, embeddings)
    print("FAISS index for Ideal Clauses created successfully.")

    # Save the index to the repository
    vectorstore.save_local(faiss_index_path)
    print(f"FAISS index saved to {faiss_index_path}.")

# Build RAG 2: General Q&A
def build_general_qa_retriever(file_path, faiss_index_path="./faiss_index_general_qa"):

    print("\n--- BUILDING RAG 2: GENERAL Q&A RETRIEVER ---")
    
    # 1. Load and process data (using the existing logic)
    try:
        df = pd.read_excel(file_path, header=1)
    except FileNotFoundError:
        print(f"Error: File '{file_path}' not found. Using dummy data.")
        df = pd.DataFrame({
            'Question': ["How do I know if my tenancy agreement is valid?"],
            'Answer / Explanation': ["Must be in writing, signed by both parties, and include essential terms."],
            'Legal Commentary': ["Only agreements containing essential terms are enforceable."],
            'Government Regulation / Explanation': ["N/A"]
        })
    
    # 2. Convert Data Rows into LangChain Documents
    documents = []
    for index, row in df.iterrows():
        content = textwrap.dedent(f"""
            Question: {row.get('Question', 'N/A')}
            Answer/Explanation: {row.get('Answer / Explanation', 'N/A')}
            Legal Context: {row.get('Legal Commentary', 'N/A')}
            Regulation/Source: {row.get('Government Regulation / Explanation', 'N/A')}
        """).strip()
        
        # Create a LangChain Document
        doc = Document(
            page_content=content,
            metadata={"source": file_path, "row_index": index}
        )
        documents.append(doc)
        
    print(f"Created {len(documents)} General Q&A Documents.")

    print("\n--- GENERAL Q&A CHUNK PREVIEW (First 3) ---")
    for i, doc in enumerate(documents[:3]):
        print(f"Chunk {i+1}:")
        print(doc.page_content)
        print(f"  Metadata: {doc.metadata}")
        print("-" * 20)

    # 3. Create Vector Store and Retriever
    vectorstore = FAISS.from_documents(documents, embeddings)
    print("FAISS index for General Q&A created successfully.")
    
    # Save the index to the repository
    vectorstore.save_local(faiss_index_path)
    print(f"FAISS index saved to {faiss_index_path}.")

def load_retriever(FAISS_INDEX_PATH, K):
    """
    Loads the saved FAISS index and returns the retriever object for runtime use.
    """
    print("\n--- LOADING RAG RETRIEVER ---")
    
    if not os.path.exists(FAISS_INDEX_PATH):
        # This is a critical error if the index should be pre-built
        raise FileNotFoundError(
            f"FAISS index not found at {FAISS_INDEX_PATH}. "
            "Please run 'create_and_save_ideal_index()' first."
        )

    # 1. Load Vector Store
    # IMPORTANT: You still need the 'embeddings' model object to load the index
    vectorstore = FAISS.load_local(
        FAISS_INDEX_PATH, 
        embeddings, 
        allow_dangerous_deserialization=True # Required by LangChain for loading
    )
    print("FAISS index loaded successfully.")

    # 2. Return Retriever
    return vectorstore.as_retriever(search_kwargs={"k": K})

def load_and_extract_pdf_text(pdf_file) -> str:
    """
    Extracts text from an IN-MEMORY PDF file (from st.file_uploader)
    and applies regex cleanup.
    
    Args:
        pdf_file: The in-memory file object (e.g., from io.BytesIO).
    
    Returns:
        A single string containing all text from the PDF.
    """
    print("--- 💬 Extracting text from in-memory PDF ---")
    
    # This function does NOT check for a file path.
    # It reads the in-memory object directly.
    
    try:
        # Use pypdf.PdfReader to read the file-like object
        reader = PdfReader(pdf_file)
        
        pages = reader.pages
        print(f"Successfully loaded {len(pages)} pages.")

        # Concatenate all page content
        full_text = "\n\n".join(page.extract_text() for page in pages if page.extract_text())
        
        # Your regex cleanup
        full_text = re.sub(r'\s{2,}', ' ', full_text)
        full_text = re.sub(r'(\n\s*){2,}', '\n\n', full_text)

        print("Text extraction complete.")
        return full_text
    
    except Exception as e:
        print(f"An error occurred during PDF text extraction: {e}")
        raise RuntimeError(f"An error occurred during PDF text extraction: {e}")

# New code for new thorough mode
class Clause(BaseModel):
    clause_title: str = Field(description="The clause number and title, e.g., '5. Maintenance' or '3. Security Deposit'")
    clause_text: str = Field(description="The full, verbatim text of the clause.")

class ParsedAgreement(BaseModel):
    clauses: List[Clause]

# Build parser chain
def get_parser_chain():
    """
    Creates and returns a runnable chain for parsing text.
    This replaces the parse_agreement_node.
    """
    print("--- 👨‍🔧 Building Parser Chain ---")
    
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    structured_llm = llm.with_structured_output(ParsedAgreement)
    
    system_prompt = """
    You are an expert legal document parser. Your sole job is to read the
    following tenancy agreement and convert it into a structured list of its
    distinct clauses.
    
    Carefully extract each clause, using its number and title (e.g., '4. Security Deposit')
    as the 'clause_title' and the full text of that clause as the 'clause_text'.
    Ensure you capture all clauses from start to finish.
    """
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        ("human", "Please parse this tenancy agreement:\n\n{agreement_text}")
    ])
    
    parser_chain = prompt | structured_llm
    return parser_chain

# Create the Chain (Run Once) 
print("--- 🚀 Building Parser Chain... ---")
parse_chain = get_parser_chain()
print("--- ✅ Parser Chain built! ---")

# Convert pdf to clauses
def process_pdf_to_clauses(uploaded_file) -> tuple[str, List[Clause]] | None:
    """
    The main processing function for PARSING ONLY.
    It handles PDF extraction and calls the simple parse_chain.
    
    Returns:
        A tuple of (raw_text, clauses), or None on failure.
    """
    
    # Step 1: Extract text from the PDF
    # (This assumes uploaded_file is an in-memory object)
    raw_text = load_and_extract_pdf_text(uploaded_file)
    
    if not raw_text or raw_text.strip() == "":
        print("--- ❌ PDF extraction failed or file is empty. ---")
        return None
        
    # Step 2: Run the compiled "Parser Chain"
    print("--- 🚀 Invoking parse_chain... ---")
    try:
        parsed_agreement = parse_chain.invoke({"agreement_text": raw_text})
        
        # Step 3: Return BOTH the raw text and the parsed clauses
        return raw_text, parsed_agreement.clauses
        
    except Exception as e:
        print(f"Error during LLM parsing: {e}")
        return None

class AnalysisGraphState(TypedDict):
    """
    The state for our analysis-only graph.
    """
    analysis_mode: str
    raw_text: str
    clauses: List[Clause] 
    retriever: Any
    checklist_path: str
    target_language: str
    missing_clauses_report: str
    clause_analyses: List[str]
    final_report: str

def setup_retriever_node(state: AnalysisGraphState) -> dict:
    """
    Loads the pre-built 'Ideal Clauses' retriever from disk.
    This is Node 2 in our graph, running after the parser.
    """
    print("--- 📚 Calling setup_retriever_node ---")
    
    try:
        # This is for the "Thorough Mode" RAG analysis
        IDEAL_CLAUSES_FAISS = os.path.join(BACKEND_SCRIPT_DIR,"faiss_index_ideal_clauses")
        ideal_retriever = load_retriever(IDEAL_CLAUSES_FAISS, K=3)
        
        # Save the retriever to the graph's state
        # The 'AnalysisGraphState' TypedDict must have a 'retriever' key
        return {"retriever": ideal_retriever}
    
    except FileNotFoundError as e:
        print(f"CRITICAL ERROR: {e}")
        # Re-raise the error to stop the graph and alert the developer
        raise e

# --- TASK A (New Node): Check for Missing Clauses ---
async def check_missing_clauses_node(state: AnalysisGraphState) -> Dict:
    """
    Uses an LLM to compare the parsed clauses against the 
    MASTER_CHECKLIST to find what's missing.
    """
    print("--- Task A: 🕵️ Checking for Missing Clauses ---")
    
    try:
        # Load the checklist
        checklist_path = state["checklist_path"]
        target_language = state.get("target_language", "English")
        checklist_items = load_checklist(checklist_path)
        master_titles = [item.get("title", "Unknown Item") for item in checklist_items]
        # Get the titles of the clauses the user *actually* has
        parsed_clause_titles = [c.clause_title for c in state["clauses"]]
        
        prompt = ChatPromptTemplate.from_template(
            """
            You are a legal analyst. Your job is to check if a tenancy 
            agreement is missing any critical clauses.
            
            Compare the "Master Checklist" (what *should* be present) 
            against the "Actual Clause Titles" (what *was* parsed 
            from the user's document).
            
            Identify and list any topics from the Master Checklist that 
            are NOT covered by the Actual Clause Titles. For each 
            missing topic, briefly state its importance.
            
            Master Checklist:
            {master_list}
            
            Actual Clause Titles:
            {actual_titles}
            
            Report on missing clauses:

            **IMPORTANT:** Translate your final answer entirely into {target_language}
            """
        )
        
        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
        
        chain = prompt | llm | StrOutputParser()
        
        report = await chain.ainvoke({
            "master_list": "\n".join(f"- {item}" for item in master_titles),
            "actual_titles": "\n".join(f"- {title}" for title in parsed_clause_titles),
            "target_language": target_language
        })
        
        return {"missing_clauses_report": report}
    except Exception as e:
        print(f"Error in missing clause check: {e}")
        return {"missing_clauses_report": "Error: Could not perform missing clause check."}

# --- TASK B (Modified Node): Run RAG on Existing Clauses ---

async def _run_rag_for_clause(clause: Clause, retriever: Any, target_language: str) -> str:
    """
    (This is the same helper function as before)
    Runs RAG to check if a single *existing* clause is "up to par".
    """
    
    rag_prompt = ChatPromptTemplate.from_template(
        """
        You are an expert tenancy agreement reviewer.
        Compare the "User's Clause" against the "Ideal Legal Context".
        
        Identify any risks, unfair terms, or deviations from the 
        ideal context. State if the clause is fair.
        
        **Ideal Legal Context:**
        {context}
        
        **User's Clause to Review:**
        {question}

        **IMPORTANT:** Translate your final analysis entirely into {target_language}.
        """
    )
    
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1)
    
    def format_docs(docs):
        return "\n\n---\n\n".join(doc.page_content for doc in docs)

    chain_setup = RunnableParallel({
        # The "context" key is populated by this sub-chain:
        # (Input -> "question" -> retriever -> format_docs)
        "context": (lambda x: x["question"]) | retriever | format_docs,
        
        # The "question" key is populated by this:
        # (Input -> "question")
        "question": (lambda x: x["question"]),
        
        # The "target_language" key is populated by this:
        # (Input -> "target_language")
        "target_language": (lambda x: x["target_language"])
    })
    
    # 3. The final chain pipes the setup into the prompt
    rag_chain = chain_setup | rag_prompt | llm | StrOutputParser()
    
    # 4. Invoke the chain with the simple dictionary
    analysis = await rag_chain.ainvoke({
        "question": clause.clause_text,
        "target_language": target_language
    })
    
    return f"### ✅ Analysis for: {clause.clause_title}\n\n{analysis}"

async def parallel_rag_node(state: AnalysisGraphState) -> Dict:
    """
    (This is also similar to before)
    Creates a parallel task for *each* existing clause.
    """
    print("--- Task B: 🔬 Running Parallel RAG Analyses ---")
    clauses = state["clauses"]
    retriever = state["retriever"]
    target_language = state.get("target_language", "English")
    
    tasks = []
    for clause in clauses:
        tasks.append(_run_rag_for_clause(clause, retriever, target_language))
    
    # Run all RAG analyses concurrently
    analyses = await asyncio.gather(*tasks)
    
    return {"clause_analyses": analyses}

# New "Hub" and "Compiler" Nodes 
async def run_parallel_analysis_hub(state: AnalysisGraphState) -> Dict:
    """
    This is the new "hub" node that runs Task A and Task B
    at the exact same time.
    """
    print("--- 🚀 Kicking off parallel analysis hub ---")
    
    # Create the two main tasks
    task_a = check_missing_clauses_node(state)
    task_b = parallel_rag_node(state)
    
    # Run them concurrently and wait for both to finish
    results = await asyncio.gather(task_a, task_b)
    
    # Merge the results from both tasks into one dictionary
    # to update the state
    combined_results = {}
    combined_results.update(results[0]) # Results from Task A
    combined_results.update(results[1]) # Results from Task B
    
    return combined_results

def compile_thorough_report_node(state: AnalysisGraphState) -> Dict:
    """
    Combines the results from both parallel tasks into one
    final markdown report for the user.
    """
    print("---  compiling Thorough Report ---")
    
    missing_report = state["missing_clauses_report"]
    analyses = state["clause_analyses"]
    
    # Format the final report
    final_report = f"""
    ## 1. Missing Clause Check
    
    {missing_report}
    
    ---
    
    ## 2. Analysis of Existing Clauses
    
    {"\n\n---\n\n".join(analyses)}
    """
    
    return {"final_report": final_report}

# Fast Mode
def compress_ta_text(raw_text: str, max_chars: int = 120_000) -> str:
    """
    Lightweight text compression: clean spaces, remove page footers,
    keep major clauses. If still too long, truncate with short summaries.
    """
    import re
    text = raw_text
    text = re.sub(r"\s+\n", "\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"-\n", "", text)
    text = re.sub(r"\n?Page \d+ of \d+\n?", "\n", text, flags=re.IGNORECASE)

    if len(text) <= max_chars:
        return text

    paras = re.split(r"\n(?=[A-Z ]{3,}\.?(\s|$)|\d{1,2}\.\s)", text)
    chunks, acc = [], 0
    for p in paras:
        p = p.strip()
        if not p:
            continue
        if acc + len(p) <= max_chars:
            chunks.append(p)
            acc += len(p)
        else:
            head = p[:1000]
            chunks.append(f"[SUMMARY] {head}")
            acc += len(head)
        if acc >= max_chars:
            break
    return "\n\n".join(chunks)


def load_checklist(checklist_path: str):
    """
    Load checklist file (JSON or CSV) and return list of dicts.
    Expected fields: id, title, requirement, keywords, must_have.
    """
    import json, os
    import pandas as pd
    ext = os.path.splitext(checklist_path)[1].lower()
    if ext == ".json":
        with open(checklist_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else data.get("items", [])
    elif ext == ".csv":
        df = pd.read_csv(checklist_path)
        return df.to_dict(orient="records")
    else:
        raise ValueError("Checklist file must be .json or .csv")


def compare_ta_with_checklist_whole(ta_text: str, checklist_items: list, model: str = None, target_language: str = "English"):
    """
    Compare entire TA (compressed) directly with checklist via LLM.
    Return structured JSON (no retrieval / embedding required).
    """
    import json, time
    global client, LLM_MODEL
    if model is None:
        model = LLM_MODEL

    if client is None:
        return {
            "items": [],
            "_error": "OpenAI client not configured. Please set OPENAI_API_KEY."
        }
    system_instruction = (
        "You are a contract compliance analyst for residential tenancy agreements. "
        "Return STRICT JSON only. For each checklist item, decide status ∈ "
        "{COMPLIANT, PARTIAL, MISSING}. Provide evidence (short TA quotes), "
        "risk (LOW|MEDIUM|HIGH) derived from the loaded checklist's criteria, "
        "recommendation (specific edit), and location_hint."
        f"**IMPORTANT:** Translate your final answer entirely into {target_language}"
    )

    user_query = f"""
=== TENANCY AGREEMENT (COMPRESSED) ===
{ta_text}

=== CHECKLIST (JSON-LIKE) ===
{checklist_items}

Return JSON with schema:
{{
  "items": [
    {{
      "id": <string|int>,
      "title": <string>,
      "status": "COMPLIANT"|"PARTIAL"|"MISSING",
      "evidence": [<short quotes>],
      "risk": "LOW"|"MEDIUM"|"HIGH",
      "recommendation": <string>,
      "location_hint": <string|null>
    }}
  ]
}}
STRICT JSON. No commentary.
""".strip()

    last_err = ""
    for _ in range(MAX_RETRIES):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_instruction},
                    {"role": "user", "content": user_query}
                ],
                response_format={"type": "json_object"},
                temperature=0.2,
            )
            content = resp.choices[0].message.content
            return json.loads(content)
        except Exception as e:
            last_err = str(e)
            time.sleep(1.5)
    return {
        "items": [],
        "_error": f"LLM comparison failed: {last_err}"
    }

def fast_checklist_node(state: AnalysisGraphState) -> dict:
    """
    Generate a report by comparing the entire TA with a checklist (no RAG).
    Automatically translates the report if target_languages is set.
    """
    raw_text = state["raw_text"]
    checklist_path = state["checklist_path"]
    target_language = state.get("target_language", "English")

    # 2a. Compress text
    max_chars = 80_000  
    ta_comp = compress_ta_text(raw_text, max_chars=max_chars)
    
    # 2b. Load checklist
    try:
        checklist = load_checklist(checklist_path)
    except Exception as e:
        print(f"Error loading checklist: {e}")
        return {"final_report": f"Error: Could not load checklist file at {checklist_path}."}

    # 2c. Compare (This calls your other helper function)
    result_json = compare_ta_with_checklist_whole(
        ta_comp, 
        checklist, 
        model=LLM_MODEL, 
        target_language=target_language
    )

    # 2d. Format into Markdown
    md_parts = []
    for item in result_json.get("items", []):
        location = item.get('location_hint','N/A')
        evidence_list = item.get('evidence',[])
        evidence_str = ("; ".join(evidence_list[:3]) or '—')
        md_parts.append(
            f"""**[{item.get('status','')}] {item.get('title','(no title)')}** - Risk: **{item.get('risk','')}**"""
        )
        if location and location != 'N/A':
            md_parts.append(f"""- Location: {location}""")
        if evidence_str != '—':
            md_parts.append(f"""- Evidence: {evidence_str}""")
        md_parts.append(
            f"""- Recommendation: {item.get('recommendation','—')}"""
        )

    final_md = "\n\n".join(md_parts) if md_parts else "No findings."

    # 3. Write the final report back to the state
    return {"final_report": final_md}

def route_analysis(state: AnalysisGraphState) -> str:
    """
    This is the "router" or "conditional edge" function.
    It reads 'analysis_mode' from the state and decides the next step.
    
    The 'analysis_mode' string ("fast" or "thorough") is passed in
    from main.py when the graph is first called.
    """
    mode = state["analysis_mode"]
    if mode == "fast":
        return "fast_mode" # must match the key
    else:
        return "thorough_mode"

print("--- 🚀 Compiling Full Analysis Graph... ---")
analysis_workflow = StateGraph(AnalysisGraphState)

# Add all nodes EXCEPT the parser
analysis_workflow.add_node("setup_retriever", setup_retriever_node)
analysis_workflow.add_node("fast_checklist", fast_checklist_node)
analysis_workflow.add_node("parallel_analysis_hub", run_parallel_analysis_hub)
analysis_workflow.add_node("compile_thorough_report", compile_thorough_report_node)

# Set the NEW entry point
analysis_workflow.set_entry_point("setup_retriever")

# Add the router (this connects setup_retriever to the two paths)
analysis_workflow.add_conditional_edges(
    "setup_retriever",
    route_analysis,
    {
        "fast_mode": "fast_checklist",
        "thorough_mode": "parallel_analysis_hub" 
    }
)

# Define the end points for each branch
analysis_workflow.add_edge("parallel_analysis_hub", "compile_thorough_report")
analysis_workflow.add_edge("compile_thorough_report", END)
analysis_workflow.add_edge("fast_checklist", END)

# Compile the final analysis app
analysis_app = analysis_workflow.compile()
print("--- ✅ Full Analysis Graph compiled! ---")

async def run_full_analysis(
    raw_text: str,
    clauses: List[Clause], 
    checklist_path: str,
    target_language: str = "English",
    analysis_mode: str = "fast"
) -> str:
    """
    The main processing function for FULL ANALYSIS.
    It takes the raw text (already parsed) and runs the full graph.
    """
    print(f"--- 🚀 Invoking full analysis_app (Mode: {analysis_mode})... ---")
    
    # The 'raw_text' will be used by the parser node again.
    # The 'analysis_mode' will be used by the router.
    inputs = {
        "raw_text": raw_text,
        "clauses": clauses, 
        "checklist_path": checklist_path,
        "target_language": target_language,
        "analysis_mode": analysis_mode
    }
    
    # We call 'analysis_app' and use 'ainvoke' for its async nodes
    final_state = await analysis_app.ainvoke(inputs)
    
    # Step 3: Return the final report
    return final_state.get('final_report')

# RAG 1: Ideal Clauses Configuration 
IDEAL_CLAUSES_FAISS = os.path.join(BACKEND_SCRIPT_DIR, "faiss_index_ideal_clauses") 
K = 3

# --- GLOBAL RESOURCE INITIALIZATION (RAG 1) ---
try:
    # Assuming 'embeddings' is globally available and load_retriever is imported
    ideal_clauses_retriever = load_retriever(IDEAL_CLAUSES_FAISS, K)
except Exception as e:
    print(f"CRITICAL: Failed to load ideal clauses RAG index: {e}")
    ideal_clauses_retriever = None

# --- RAG 2: General Q&A Configuration ---
GENERAL_QA_FAISS = os.path.join(BACKEND_SCRIPT_DIR,"faiss_index_general_qa" )

# --- GLOBAL RESOURCE INITIALIZATION (RAG 2) ---
try:
    # Assuming 'embeddings' is globally available and load_retriever is imported
    general_qa_retriever = load_retriever(GENERAL_QA_FAISS, K)
except Exception as e:
    print(f"CRITICAL: Failed to initialize global RAG components: {e}")
    general_qa_retriever = None

# Translation feature
def translate_document(file_path: str, target_language: list) -> dict:
    """
    Translate the content of a PDF or DOCX document from a file path using OpenAI LLM.

    Args:
        file_path: The file path to the uploaded document
        target_languages: List of target languages, e.g., ["Indonesian", "Mandarin"]

    Returns:
        dict: language -> translated text, or {"error": "message"}
    """
    if not target_language or not isinstance(target_language, list):
        return {"error": "target_languages must be a non-empty list of languages."}

    try:
        text = ""
        
        # --- ADD LOGIC TO READ THE FILE FROM THE PATH ---
        if file_path.lower().endswith(".pdf"):
            # Open the file path in read-binary mode
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join([p.extract_text() or "" for p in reader.pages])
        
        elif file_path.lower().endswith(".docx"):
            # Use the docx library to read the file path
            document = docx.Document(file_path)
            text = "\n".join([para.text for para in document.paragraphs])
            
        else:
            return {"error": "Unsupported file type. Only .pdf and .docx are supported."}
        # --- END OF FILE READING LOGIC ---

        if not text.strip():
            # This catches scanned PDFs (image-based)
            return {"error": "File is empty or text could not be extracted (check if PDF is scan/image)."}

        # Initialize LLM
        llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)

        translations = {}
        for lang in target_language:
            prompt = f"Translate the following text into {lang}:\n\n{text}"
            
            # Use .invoke() - this is the more modern way in LangChain
            response = llm.invoke(prompt) 
            translations[lang] = response.content

        return translations

    except Exception as e:
        # A more general error message
        return {"error": f"Failed to process document: {str(e)}"}

# Initialisation function for session state
def initialize_qa_resources(openai_api_key: str):
    global client
    """Initializes the OpenAI client (runtime dependency)."""
    
    # 1. Initialize OpenAI Client
    try:
        from openai import OpenAI
        client = OpenAI(api_key=openai_api_key)
    except Exception as e:
        raise RuntimeError(f"OpenAI client configuration failed: {e}")

    # 2. Return the client and the globally initialized retriever
    if general_qa_retriever is None:
        raise RuntimeError("General QA Retriever failed to load during module import.")
        
    return client, general_qa_retriever

# Integration of LLM and RAG 2
def answer_contextual_question_openai(
    user_question: str, 
    general_qa_retriever: object,
    ta_report: Optional[List[Dict]] = None, # Optional TA Report (Phase 1 output)
    past_messages: Optional[List[Dict[str, str]]] = None, # Optional Chat History
    target_language: str = "English"
) -> str:
    """
    Answers a user question using RAG 2, the TA report, and chat history.
    """
    if not client:
        return "OpenAI client is not configured correctly. Check API key."

    print(f"\n[Q&A] Answering contextual question with GPT: {user_question[:50]}...")

    # 1. Use RAG 2 to retrieve the General Q&A context
    qa_context = general_qa_retriever.invoke(user_question)
    
    # 2. Format Contexts
    general_context_str = "\n---\n".join([doc.page_content for doc in qa_context])
    report_context_str = str(ta_report) if ta_report else "No specific document analysis report available." # Format report if provided

    # 3. DEFINE THE LLM MESSAGE LIST (Chat History Integration)
    
    # Update System Instruction to acknowledge all contexts
    system_instruction = (
        "You are an expert legal assistant specializing in residential tenancy agreements. "
        "Your response must be based on the provided context, which includes the **General Law Context** (RAG data) "
        "and, if present, the **User Document Analysis Report** (specific critiques). "
        "You are capable of understanding questions in various languages, but your primary source material is English. "
        "Maintain the flow of the conversation history. If the user's question relates to a flagged clause, "
        "prioritize the information from the Analysis Report. "
        "Be concise and professional."
        f"**IMPORTANT:** Translate your answer entirely into {target_language}"
    )
    
    messages = [
        {"role": "system", "content": system_instruction}
    ]

    # Add past conversation messages if they exist
    if past_messages:
        messages.extend(past_messages) # Extends the list with previous turns

    # 4. Construct the Final User Query (containing all RAG context)
    final_user_prompt = textwrap.dedent(f"""
        Please answer the FINAL USER QUESTION using the context provided below.

        **THIS IS A REPORT ON THE USER'S TENANCY AGREEMENT (Specific Critiques):**
        ---
        {report_context_str}
        
        **GENERAL Q&A CONTEXT (RAG Retrieved):**
        ---
        {general_context_str}
        ---

        **FINAL USER QUESTION:** {user_question}
    """)
    
    # Add the final, context-rich user prompt
    messages.append({"role": "user", "content": final_user_prompt})

    # 5. EXECUTE API CALL WITH EXPONENTIAL BACKOFF
    for attempt in range(MAX_RETRIES):
        try:
         
            response = client.chat.completions.create(
                model=LLM_MODEL,
                messages=messages, # Now uses the full history list
                temperature=0.0
            )
            answer_text = response.choices[0].message.content
            styled_answer = f"""<span style="color:black;">{answer_text}</span>""" 
            
            print(f"... Successful answer generated on attempt {attempt + 1}.")
            return styled_answer
            
        except Exception as e:
            if attempt < MAX_RETRIES - 1:
                wait_time = 2 ** attempt
                print(f"OpenAI API Error: {e}. Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                print(f"OpenAI API failed after {MAX_RETRIES} attempts.")
                return f"Unable to generate answer after {MAX_RETRIES} attempts. Error: {e}"
    
    return "Unknown error."